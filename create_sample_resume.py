#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt

doc = Document()

# Add heading
heading = doc.add_paragraph()
heading_run = heading.add_run('ALEX MORGAN')
heading_run.bold = True
heading_run.font.size = Pt(14)
heading.alignment = 1  # Center

# Contact info
contact = doc.add_paragraph('Email: alex.morgan@email.com | Phone: (555) 123-4567 | LinkedIn: linkedin.com/in/alexmorgan | GitHub: github.com/alexmorgan')
contact.alignment = 1

# Summary
doc.add_heading('PROFESSIONAL SUMMARY', level=2)
doc.add_paragraph('Experienced software engineer with 5+ years in full-stack development. Proficient in Python, JavaScript, React, and cloud platforms. Proven track record of building scalable applications and leading technical initiatives.')

# Experience
doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)

doc.add_heading('Senior Software Engineer', level=3)
doc.add_paragraph('Tech Company Inc. | Jan 2022 - Present')
doc.add_paragraph('Led development of microservices architecture using Python and Docker')
for item in ['Reduced API response time by 40% through optimization', 'Mentored 3 junior developers on best practices', 'Implemented CI/CD pipeline using GitHub Actions']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Full Stack Developer', level=3)
doc.add_paragraph('StartUp XYZ | June 2019 - Dec 2021')
doc.add_paragraph('Developed and maintained full-stack web applications using React and Node.js')
for item in ['Built real-time dashboard used by 10,000+ users', 'Implemented automated testing reducing bugs by 35%', 'Collaborated with product team to deliver features on schedule']:
    doc.add_paragraph(item, style='List Bullet')

# Skills
doc.add_heading('TECHNICAL SKILLS', level=2)
skills = 'Languages: Python, JavaScript, TypeScript, SQL | Frameworks: React, Node.js, Django, FastAPI | Tools: Docker, Kubernetes, AWS, Git | Databases: PostgreSQL, MongoDB, Redis'
doc.add_paragraph(skills)

# Education
doc.add_heading('EDUCATION', level=2)
doc.add_paragraph('Bachelor of Science in Computer Science')
doc.add_paragraph('State University | 2019')

doc.save('resumes/candidate_resume.docx')
print('Resume created successfully at resumes/candidate_resume.docx')
